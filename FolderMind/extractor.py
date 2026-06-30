import os

# Try to import PDF libraries gracefully
try:
    import pypdf
except ImportError:
    try:
        import PyPDF2 as pypdf
    except ImportError:
        pypdf = None

# Try to import python-docx gracefully
try:
    import docx
except ImportError:
    docx = None

def extract_text_from_file(filepath):
    """
    Extracts text from txt, md, pdf, or docx files.
    Truncates text to ~8000 characters to prevent context overflow.
    """
    ext = os.path.splitext(filepath)[1].lower()
    text = ""

    if ext in ('.txt', '.md', '.json', '.csv', '.py', '.js', '.html', '.css'):
        text = extract_txt(filepath)
    elif ext == '.pdf':
        text = extract_pdf(filepath)
    elif ext in ('.docx', '.doc'):
        text = extract_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    # Clean up and truncate
    text = text.strip()
    if len(text) > 8000:
        text = text[:8000] + "\n\n[... TEXT TRUNCATED BY FOLDERMIND TO AVOID AI CONTEXT OVERFLOW ...]"
    
    return text

def extract_txt(filepath):
    """Reads a plain text file with robust encoding fallback."""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
    for encoding in encodings:
        try:
            with open(filepath, 'r', encoding=encoding, errors='ignore') as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode plain text file with standard encodings.")

def extract_pdf(filepath):
    """Extracts text from a PDF file."""
    if pypdf is None:
        raise ImportError("No PDF reader library installed (pypdf or PyPDF2). Please install pypdf.")
    
    text_parts = []
    try:
        with open(filepath, 'rb') as f:
            reader = pypdf.PdfReader(f)
            # Read first 15 pages maximum (to keep within reasonable processing limit)
            max_pages = min(15, len(reader.pages))
            for i in range(max_pages):
                page_text = reader.pages[i].extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        raise RuntimeError(f"Error reading PDF file: {str(e)}")
        
    return "\n\n".join(text_parts)

def extract_docx(filepath):
    """Extracts text from a DOCX file."""
    if docx is None:
        raise ImportError("python-docx library is not installed. Please install python-docx.")
    
    text_parts = []
    try:
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)
    except Exception as e:
        raise RuntimeError(f"Error reading DOCX file: {str(e)}")
        
    return "\n".join(text_parts)
